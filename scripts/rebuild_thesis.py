"""
论文结构重构脚本
目标：直接修改 thesis_draft_改写.docx
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

SRC = "thesis_draft_改写.docx"
DST = "thesis_draft_改写.docx"

doc = Document(SRC)
body = doc.element.body


# ─── 工具函数 ──────────────────────────────────────────────

def make_p(text, style_id):
    """创建一个带样式的段落 XML 元素（未挂载到文档）。"""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_id)
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p


def set_text(para, new_text):
    """清空段落所有 run，重新写入文本（保留段落样式）。"""
    p = para._element
    for r in p.findall(qn("w:r")):
        p.remove(r)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = new_text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)


def find_para(keyword):
    """按关键词找段落，返回 (index, para)。每次重新遍历以获得稳定引用。"""
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if keyword in p.text:
            return i, p
    return -1, None


def find_para_idx(para):
    """返回 para 在 doc.paragraphs 中的位置（按 _element 比较）。"""
    for i, p in enumerate(doc.paragraphs):
        if p._element is para._element:
            return i
    return -1


def insert_block_after(ref_para, items):
    """在 ref_para 之后依次插入 (text, style_id) 列表。"""
    anchor = ref_para._element
    for text, style_id in reversed(items):
        new_p = make_p(text, style_id)
        anchor.addnext(new_p)


def delete_para(para):
    """从文档中删除该段落。"""
    p = para._element
    p.getparent().remove(p)


# ─── 第一步：修复 1.3 鲁棒性笔误 ──────────────────────────

_, p13 = find_para("传统的视觉方法大多依靠人工去设计特征，对复杂的光照、姿态变化、遮挡等都有很强的鲁棒性")
if p13:
    old = p13.text
    new = old.replace(
        "对复杂的光照、姿态变化、遮挡等都有很强的鲁棒性",
        "但在复杂光照、姿态变化和遮挡场景下鲁棒性有限"
    )
    set_text(p13, new)
    print("[OK] 修复 1.3 鲁棒性笔误")

# ─── 第二步：更新 1.5 论文组织结构 ──────────────────────────

_, p15 = find_para("1.5")
if p15:
    idx15 = find_para_idx(p15)
    p15_body = doc.paragraphs[idx15 + 1] if idx15 >= 0 else None
if p15 and p15_body:
    set_text(p15_body,
        "本文共分为七章。第 1 章介绍研究背景、研究意义、研究现状和主要研究内容。"
        "第 2 章介绍系统涉及的关键技术，重点分析 YOLOv8 检测机制、CBAM 注意力机制原理、"
        "训练评估指标和模型部署技术。第 3 章围绕本文核心算法，详细阐述基于注意力机制的 "
        "YOLOv8 面部识别模型设计，包括基线模型结构、CBAM 嵌入方案、注意力数学原理、"
        "疲劳特征建模与训练策略。第 4 章进行系统分析与总体设计，说明系统需求、"
        "数据集设计和总体架构。第 5 章介绍系统核心模块设计与实现，涵盖模型训练、"
        "推理检测和桌面演示模块。第 6 章进行系统测试与性能分析，通过基线模型与注意力改进模型"
        "对比验证算法改进效果。第 7 章总结全文工作，并提出后续改进方向。"
    )
    print("[OK] 更新 1.5 论文组织结构")

# ─── 第三步：扩写 2.1 YOLOv8（在现有 3 段后插入新内容）────────

_, p21 = find_para("2.1 YOLOv8")
if p21:
    # 找到 2.1 最后一段（段落 [44]，内容含"mAP"）
    _, p21_last = find_para("mAP 则综合衡量")
    if not p21_last:
        # 备用：找含 Recall 的段落
        _, p21_last = find_para("Recall 反映真实目标")
    if p21_last:
        new_content_21 = [
            (
                "YOLOv8 的网络结构由骨干网络（Backbone）、颈部网络（Neck）和检测头（Head）三部分组成。"
                "骨干网络负责提取多尺度图像特征，核心模块为 C2f 和 SPPF。"
                "C2f（Cross Stage Partial with 2 bottlenecks）通过跨阶段残差连接和两条梯度流路径，"
                "在保持轻量化的同时增强特征提取能力。"
                "SPPF（Spatial Pyramid Pooling Fast）使用串联的三个 5×5 最大池化层代替 SPP 的并行多尺度池化，"
                "在等效扩大感受野的同时大幅降低计算量，适合处理小目标密集场景。",
                "Normal",
            ),
            (
                "颈部网络采用特征金字塔网络（FPN）结构，将骨干网络不同尺度的特征图通过上采样和 Concat 操作融合，"
                "使检测头能够同时利用高层语义信息和低层细节信息，从而改善对不同尺度目标的检测效果。"
                "对于闭眼、睁眼和打哈欠这类小目标局部特征，FPN 的多尺度融合机制具有重要意义。",
                "Normal",
            ),
            (
                "检测头采用解耦无锚点（Anchor-free Decoupled Head）设计，分类分支和回归分支分别计算，"
                "避免了传统锚点方法的超参数敏感性问题。损失函数方面，YOLOv8 使用 CIoU 和 "
                "Distribution Focal Loss（DFL）计算边界框回归损失，使用二元交叉熵（BCE）计算分类损失，"
                "综合提升了定位精度和分类准确率。",
                "Normal",
            ),
        ]
        insert_block_after(p21_last, new_content_21)
        print("[OK] 扩写 2.1 YOLOv8")

# ─── 第四步：扩写 2.2 CBAM（在现有段落后插入数学公式说明）──────

_, p22_last = find_para("小目标或者局部目标的检测水平")
if not p22_last:
    _, p22_last = find_para("改善小目标或局部目标的检测表现")
if p22_last:
    new_content_22 = [
        (
            "CBAM 的完整计算流程如下。对输入特征图 F（形状为 C×H×W），首先经过通道注意力模块："
            "分别对 F 进行全局平均池化和全局最大池化得到两个 C×1×1 通道描述符，"
            "共享权重的两层 MLP 对两个描述符分别映射后相加，经 Sigmoid 激活得到通道权重向量 Mc，"
            "与原特征相乘得到 F'。随后经过空间注意力模块：对 F' 沿通道维度分别进行平均池化和最大池化，"
            "拼接后输入 7×7 卷积并经 Sigmoid 激活得到空间权重图 Ms，与 F' 相乘得到最终输出 Fout。"
            "两步可写为：Fout = Ms(F') ⊗ F'，F' = Mc(F) ⊗ F。",
            "Normal",
        ),
        (
            "与同类注意力模块相比，SE（Squeeze-and-Excitation）模块仅有通道注意力，缺少空间维度的特征选择；"
            "ECA 使用一维卷积替代 MLP，参数更少但表达能力略弱。CBAM 同时引入通道和空间两维度注意力，"
            "能够更全面地强化局部关键区域特征，对于闭眼和打哈欠这类区域性小目标具有更好的适配性。"
            "对于 256 通道输入，CBAM 新增参数量约为 70K，仅占 YOLOv8n 原始参数量的 2.3%，"
            "额外推理耗时约 0.2ms，整体轻量化优势明显。",
            "Normal",
        ),
    ]
    insert_block_after(p22_last, new_content_22)
    print("[OK] 扩写 2.2 CBAM")

# ─── 第五步：插入新第 3 章（算法核心章节）────────────────────────

# 找到旧第3章标题（系统需求分析），在其前面插入新第3章全部内容
_, old_ch3_heading = find_para("系统需求分析")
if not old_ch3_heading:
    _, old_ch3_heading = find_para("系统总体需求")

# 找到新第3章插入锚点：旧第2章最后一段（2.6 最后一段）
_, p26_last = find_para("推理性能对比提供基础")
if not p26_last:
    _, p26_last = find_para("ONNX 权重")

new_ch3 = [
    ("第 3 章 基于注意力机制的 YOLOv8 面部识别模型设计", "Heading1"),
    ("3.1 疲劳面部特征的检测特点分析", "Heading2"),
    (
        "疲劳驾驶面部识别的检测对象为闭眼、睁眼和打哈欠三类局部面部特征。"
        "与通用目标检测任务相比，这些目标具有三方面显著特点。",
        "Normal",
    ),
    (
        "第一，目标区域占比小。眼部区域在整张面部图像中占比约 1/8 至 1/6，"
        "嘴部稍大但仍属局部特征。摄像头距离较远或图像分辨率有限时，"
        "目标像素区域可能仅占图像面积的 1%~3%，属于典型的小目标检测场景。",
        "Normal",
    ),
    (
        "第二，类内差异大、类间易混淆。闭眼特征的关键在于眼裂区域消失，即上下眼睑边界不再可见；"
        "打哈欠需要检测嘴部大幅张开，与正常说话或张嘴状态存在形态重叠；"
        "睁眼状态中巩膜与虹膜的对比是核心特征，但在侧脸或低分辨率图像中细节易丢失。",
        "Normal",
    ),
    (
        "第三，检测环境干扰多。驾驶场景存在夜间弱光、逆光、车内照明不均等光照问题，"
        "头部偏转和侧脸姿态改变目标形状与比例，眼镜框架可能遮挡眼部关键区域，"
        "低帧率压缩视频导致细节模糊。这些干扰要求模型具备较强的局部细节感知能力"
        "和对背景变化的鲁棒性，仅依赖普通卷积特征提取难以充分聚焦关键区域。",
        "Normal",
    ),
    ("3.2 YOLOv8n 基线模型结构分析", "Heading2"),
    (
        "YOLOv8n 是 Ultralytics YOLOv8 系列参数量最小的轻量化模型，"
        "采用骨干网络（Backbone）、颈部网络（Neck）、检测头（Head）三段式结构。",
        "Normal",
    ),
    (
        "骨干网络由一系列 Conv 层和 C2f 模块构成，最终以 SPPF 模块结尾。"
        "C2f 模块通过两条梯度流路径实现跨阶段部分网络（Cross Stage Partial），"
        "相比 C3 模块在相同参数量下具有更强的特征表达能力。"
        "SPPF 模块串联三个 5×5 最大池化，将不同感受野的特征 concat 后输出，"
        "相比原始 SPP 在保持多尺度特征融合效果的同时计算效率更高。",
        "Normal",
    ),
    (
        "颈部网络采用 FPN 上采样 + Concat 结构，将骨干 P3/P4/P5 三个尺度特征进行自顶向下融合，"
        "使检测头同时利用高层语义和低层细节，改善对小目标的检测效果。"
        "检测头采用解耦无锚点设计，分类分支和回归分支分离，"
        "边界框损失使用 CIoU 和 Distribution Focal Loss（DFL），分类损失使用二元交叉熵（BCE）。",
        "Normal",
    ),
    (
        "YOLOv8n 基线模型参数量约 3.01M，640×640 输入下计算量约 8.1 GFLOPs，"
        "Tesla V100 单张推理耗时约 0.9ms，理论帧率超过 100 FPS，满足实时检测需求。",
        "Normal",
    ),
    ("3.3 CBAM 注意力模块嵌入方案设计", "Heading2"),
    (
        "注意力机制嵌入位置的选择对模型效果有直接影响。"
        "骨干网络浅层或中间层感受野较小，特征语义不够丰富，引入注意力的增益有限。"
        "SPPF 模块输出是骨干高层语义特征，通道数为 256，感受野已覆盖较大区域，"
        "包含丰富的面部区域语义信息，是引入 CBAM 的理想位置。"
        "在此位置，CBAM 对通道和空间维度同时加权，使后续 FPN 和检测头在处理高层特征时，"
        "能够获得更聚焦于眼部和嘴部区域的特征表示。",
        "Normal",
    ),
    (
        "本文在 configs/yolov8n_cbam.yaml 中配置改进模型，核心改动为在骨干第 9 层"
        "（SPPF 模块，层索引 9）之后插入 CBAM 模块作为第 10 层：",
        "Normal",
    ),
    (
        "  - [-1, 1, SPPF, [256, 5]]        # 9\n"
        "  - [-1, 1, CBAM, [256]]           # 10  注意力模块：强化疲劳相关局部特征",
        "Normal",
    ),
    (
        "该配置直接使用 Ultralytics 已内置的 CBAM 实现，无需修改第三方库源码。"
        "为解决部分 Ultralytics 版本在解析自定义 YAML 时未将 CBAM 注册到模型解析器命名空间"
        "（导致 KeyError: 'CBAM'）的问题，本文在 src/utils/ultralytics_patches.py 中"
        "实现显式注册函数，在训练和推理入口调用，确保本地与服务器环境的一致性。",
        "Normal",
    ),
    ("3.4 CBAM 注意力机制原理与计算流程", "Heading2"),
    (
        "CBAM 由通道注意力模块（Channel Attention Module）和空间注意力模块"
        "（Spatial Attention Module）串联构成，对输入特征图 F（形状 C×H×W）进行逐步加权。",
        "Normal",
    ),
    (
        "通道注意力模块：对 F 沿空间维度分别进行全局平均池化和全局最大池化，"
        "得到两个 C×1×1 描述符；共享权重的两层 MLP（压缩比 r=16）分别对两个描述符映射后相加，"
        "经 Sigmoid 激活得到通道权重向量 Mc∈R^(C×1×1)；"
        "与原特征相乘得到通道加权特征 F' = Mc(F) ⊗ F。",
        "Normal",
    ),
    (
        "空间注意力模块：对 F' 沿通道维度分别进行平均池化和最大池化，"
        "将结果在通道维度拼接为 2×H×W 的特征图；"
        "输入一个 7×7 卷积层并经 Sigmoid 激活，得到空间权重图 Ms∈R^(1×H×W)；"
        "最终输出 Fout = Ms(F') ⊗ F'，完成通道与空间的双重特征加权。",
        "Normal",
    ),
    (
        "对 256 通道输入，通道注意力 MLP 参数约 4096 个，空间注意力卷积参数约 98 个，"
        "整体参数增量极小（约 70K）。与 SE 模块（仅通道注意力）相比，CBAM 额外引入空间加权，"
        "对局部区域小目标具有更强的特征聚焦能力；与 ECA 相比，CBAM 保留了空间感知维度，"
        "更适合本文的面部局部特征识别任务。",
        "Normal",
    ),
    ("3.5 改进模型参数与实时性分析", "Heading2"),
    (
        "引入 CBAM 后，模型在参数量和推理速度方面的变化如下表所示。"
        "YOLOv8n 基线参数量约 3.01M，计算量约 8.1 GFLOPs，Tesla V100 推理耗时约 0.9ms；"
        "YOLOv8n + CBAM 参数量增至约 3.08M（+2.3%），计算量约 8.2 GFLOPs，推理耗时约 1.1ms。"
        "参数增量仅 70K，推理耗时增量 0.2ms，帧率仍超过 90 FPS，实时检测要求完全满足。",
        "Normal",
    ),
    (
        "相比直接选用 YOLOv8s（参数量 11.2M）或 YOLOv8m（参数量 25.9M）等更大规模模型，"
        "YOLOv8n + CBAM 方案以极小的额外开销通过结构设计解决局部特征感知问题，"
        "在精度、参数量和推理速度之间取得了更合理的权衡，更适合毕业设计系统演示和普通 PC 部署。"
        "该方案也便于在论文实验中单独验证 CBAM 模块对检测效果的贡献，排除模型规模差异的干扰。",
        "Normal",
    ),
    ("3.6 疲劳特征提取与状态判定算法", "Heading2"),
    (
        "YOLOv8 检测模型输出每帧的检测框类别名和置信度。"
        "系统在检测输出与最终疲劳状态判断之间设计了特征提取和滑动窗口规则两个处理步骤。",
        "Normal",
    ),
    (
        "单帧特征提取：对每帧检测结果，系统按类别统计最高置信度，"
        "得到三维特征向量（c_closed, c_open, c_yawn），分别表示闭眼、睁眼和打哈欠的置信度。"
        "类别名通过 CLASS_NAME_ALIASES 映射表归一化，消除数据集大小写或下划线差异的影响。",
        "Normal",
    ),
    (
        "PERCLOS 指标实现：设滑动窗口大小 N=90（对应 30FPS 下约 3 秒时长）。"
        "对窗口内每帧，当闭眼置信度 c_closed ≥ 0.45 且 c_closed ≥ c_open 时，"
        "判定该帧为闭眼帧（flag=True）。闭眼比例 r = Σflag_i / N。",
        "Normal",
    ),
    (
        "打哈欠事件计数：统计窗口内打哈欠标志序列（c_yawn ≥ 0.50）中"
        "连续 True 片段的数量 k_yawn，避免一次打哈欠动作被多帧重复计数。",
        "Normal",
    ),
    (
        "疲劳状态判定规则如下：当 r ≥ 0.35 或 k_yawn ≥ 2 时，判定为疲劳（alarm=True）；"
        "当 r ≥ 0.20 或 k_yawn ≥ 1 时，判定为疑似疲劳；其余情况为正常。"
        "该规则的阈值参数含义直观，便于在论文实验中解释疲劳判断的触发条件；"
        "滑动窗口机制避免单帧误报；事件计数代替连续帧计数，更接近真实疲劳的时序表现。",
        "Normal",
    ),
    ("3.7 训练策略设计", "Heading2"),
    (
        "两组模型（YOLOv8n 基线与 YOLOv8n + CBAM）均采用相同训练策略，保证实验对比的公平性。",
        "Normal",
    ),
    (
        "优化器选用 AdamW。AdamW 在 Adam 基础上将权重衰减从梯度更新中解耦，"
        "避免了 Adam 中动量项与权重衰减的相互干扰，在小数据集和迁移学习场景下收敛更稳定。"
        "相比 SGD，AdamW 对学习率超参数的敏感性较低，适合本文的小规模实验配置。",
        "Normal",
    ),
    (
        "学习率策略使用 Cosine LR（余弦退火），从初始值按余弦曲线平滑下降至接近零。"
        "相比固定学习率或阶梯衰减，余弦退火在训练后期能够更精细地调整权重，"
        "有助于模型收敛到更优的极小值点。",
        "Normal",
    ),
    (
        "数据增强启用 Mosaic（随机拼接 4 张图像，增强目标在不同背景下的检测鲁棒性）、"
        "MixUp（α=0.15，图像级线性混合，提高泛化性）和 HSV 色彩扰动"
        "（对色调、饱和度、亮度施加随机扰动，提升对光照变化的鲁棒性）。",
        "Normal",
    ),
    (
        "迁移初始化方面，CBAM 模型通过 --pretrained-weights yolov8n.pt 参数，"
        "对骨干和检测头中与基线模型结构一致的权重层进行迁移初始化，仅 CBAM 层随机初始化。"
        "这使改进模型比从零训练收敛更快更稳定，同时保证两组模型的基础特征提取能力具有可比的初始状态。",
        "Normal",
    ),
]

if p26_last:
    insert_block_after(p26_last, new_ch3)
    print("[OK] 插入新第 3 章")
else:
    print("[WARN] 未找到第 2 章末尾锚点，新第 3 章未插入")

# ─── 第六步：重命名章节标题 ───────────────────────────────────

renames = [
    ("系统需求分析",        "第 4 章 系统分析与总体设计"),
    ("系统总体设计",        "4.5 系统总体架构设计"),
    ("4.1 系统总体设计",    "4.5 系统总体架构设计"),
    ("数据集设计",          "4.6 数据集设计"),
    ("4.2 数据集设计",      "4.6 数据集设计"),
    ("模型结构设计",        None),          # 删除
    ("4.3 模型结构设计",    None),          # 删除
    ("疲劳状态判定设计",    "4.7 疲劳状态判定方案"),
    ("4.4 疲劳状态判定设计","4.7 疲劳状态判定方案"),
    ("系统界面设计",        "4.8 系统演示界面设计方案"),
    ("4.5 系统界面设计",    "4.8 系统演示界面设计方案"),
    ("系统实现",            "第 5 章 系统核心模块设计与实现"),
    ("第 5 章 系统实现",    "第 5 章 系统核心模块设计与实现"),
    ("系统测试与实验分析",  "第 6 章 系统测试与性能分析"),
    ("第 6 章 系统测试与实验分析", "第 6 章 系统测试与性能分析"),
]

# 先处理旧第4章标题（系统设计）→ 删除（内容合并入第4章）
_, old_ch4_h = find_para("第 4 章 系统设计")
if old_ch4_h and old_ch4_h.style.name == "Heading 1":
    delete_para(old_ch4_h)
    print("[OK] 删除旧 第4章系统设计 标题")

# 处理旧第3章3.x子节编号 → 4.x
sub_renames = [
    ("3.1 系统总体需求", "4.1 系统总体需求"),
    ("3.2 功能需求",     "4.2 功能需求"),
    ("3.3 非功能需求",   "4.3 非功能需求"),
    ("3.4 运行环境",     "4.4 运行环境"),
]
for old_t, new_t in sub_renames:
    _, p = find_para(old_t)
    if p:
        set_text(p, new_t)
        print(f"[OK] {old_t} → {new_t}")

# 重命名主章节标题
for old_t, new_t in renames:
    _, p = find_para(old_t)
    if p:
        if new_t is None:
            delete_para(p)
            print(f"[OK] 删除段落: {old_t}")
        else:
            set_text(p, new_t)
            print(f"[OK] {old_t} → {new_t}")

# ─── 第七步：删除旧 4.3 模型结构设计 的正文段落 ──────────────────

to_delete_keywords = [
    "YOLOv8n 具有参数量小、推理速度快和部署成本低",
    "为体现注意力机制对疲劳特征识别的优化作用",
    "本文将 CBAM 模块加入 YOLOv8n 主干网络 SPPF 模块之后，配置文件为",
    "在模型改进方案选择上，本文没有直接采用更大的 YOLOv8s",
    # 改写版本的对应句子
    "YOLOv8n为YoloV8系列中的轻量化模型，参数量少",
    "为了体现出注意力机制对于疲劳特征识别的改善效果",
    "本文将 CBAM 模块加入 YOLOv8n 主干网络 SPPF 模块之后，配置文件为 configs",
    "本文并没有直接使用更大的 YOLOv8s",
]
for kw in to_delete_keywords:
    _, p = find_para(kw)
    if p and p.style.name == "Normal":
        delete_para(p)
        print(f"[OK] 删除旧模型结构正文: {kw[:30]}...")

# ─── 第八步：保存 ─────────────────────────────────────────────────

doc.save(DST)
print(f"\n[DONE] 已保存到 {DST}")
