"""将论文 Markdown 主稿转换为 Word 草稿。

本脚本用于毕业论文材料整理：从 docs/thesis_draft.md 读取正文，
生成 docs/thesis_draft.docx，并尽量保留标题、表格、图片和图题。
脚本依赖 python-docx，推荐使用 Codex 文档运行时执行；项目 .venv
若安装了 python-docx 也可以直接运行。
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = Path(__file__).resolve().parent
INPUT_MD = DOCS_DIR / "thesis_draft.md"
OUTPUT_DOCX = DOCS_DIR / "thesis_draft.docx"


def set_east_asia_font(run, font_name: str = "宋体") -> None:
    """为中文内容设置东亚字体，避免 Word 中出现中英文字体混乱。"""

    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    """设置表格单元格文字，并统一中文字体和字号。"""

    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text.strip())
    run.bold = bold
    run.font.size = Pt(10.5)
    set_east_asia_font(run)


def set_table_borders(table) -> None:
    """为论文表格设置细边框，使渲染后更接近正式文档样式。"""

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def parse_table(lines: list[str], start_index: int) -> tuple[list[list[str]], int]:
    """解析 Markdown 表格，返回二维内容和表格结束后的行号。"""

    rows: list[list[str]] = []
    index = start_index
    while index < len(lines) and lines[index].strip().startswith("|"):
        line = lines[index].strip().strip("|")
        cells = [cell.strip() for cell in line.split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def add_markdown_text(paragraph, text: str) -> None:
    """写入一段普通文本，并处理少量 Markdown 行内标记。"""

    # 论文正文中反引号主要用于路径和命令，Word 里保留原文即可。
    cleaned = text.replace("`", "")
    run = paragraph.add_run(cleaned)
    run.font.size = Pt(10.5)
    set_east_asia_font(run)


def add_picture(document: Document, image_ref: str, caption: str) -> None:
    """插入图片和居中图题，图片宽度控制在页面可读范围内。"""

    image_path = (DOCS_DIR / image_ref).resolve()
    if not image_path.exists():
        image_path = (ROOT / image_ref).resolve()
    if not image_path.exists():
        paragraph = document.add_paragraph()
        add_markdown_text(paragraph, f"[图片缺失：{image_ref}]")
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    # 实验曲线图以横向阅读为主，按宽度控制；竖向截图若只按宽度缩放，
    # 高度会过大并把图题挤到下一页，因此根据原图宽高比选择更合适的尺寸。
    normalized_ref = image_ref.replace("\\", "/")
    with Image.open(image_path) as image:
        width_px, height_px = image.size
    aspect = width_px / max(height_px, 1)
    if "runs" in normalized_ref:
        run.add_picture(str(image_path), width=Cm(12.8))
    elif width_px < 300 or height_px < 300:
        run.add_picture(str(image_path), width=Cm(6.2))
    elif aspect < 0.75:
        run.add_picture(str(image_path), height=Cm(13.2))
    else:
        run.add_picture(str(image_path), width=Cm(10.5))

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    caption_run.font.size = Pt(9.5)
    set_east_asia_font(caption_run)


def configure_document(document: Document) -> None:
    """设置页面、正文样式和标题样式。"""

    section = document.sections[0]
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(0.74)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def build_docx() -> None:
    """执行 Markdown 到 DOCX 的转换。"""

    text = INPUT_MD.read_text(encoding="utf-8")
    lines = text.splitlines()

    document = Document()
    configure_document(document)

    index = 0
    while index < len(lines):
        line = lines[index].rstrip()

        if not line:
            index += 1
            continue

        if line.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(18)
            set_east_asia_font(run, "黑体")
            index += 1
            continue

        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
            index += 1
            continue

        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
            index += 1
            continue

        image_match = re.match(r"!\[(.+?)\]\((.+?)\)", line)
        if image_match:
            caption, image_ref = image_match.groups()
            add_picture(document, image_ref, caption)
            index += 1
            continue

        if line.startswith("|"):
            rows, next_index = parse_table(lines, index)
            if rows:
                table = document.add_table(rows=len(rows), cols=len(rows[0]))
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_table_borders(table)
                for row_index, row in enumerate(rows):
                    for col_index, cell_text in enumerate(row):
                        set_cell_text(table.cell(row_index, col_index), cell_text, bold=row_index == 0)
                document.add_paragraph()
            index = next_index
            continue

        if line.startswith("```"):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index])
                index += 1
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Cm(0)
            run = paragraph.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            index += 1
            continue

        if line.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Cm(0)
            run = paragraph.add_run(line[2:].strip())
            run.italic = True
            run.font.size = Pt(10.5)
            set_east_asia_font(run, "宋体")
            index += 1
            continue

        paragraph = document.add_paragraph()
        add_markdown_text(paragraph, line)
        index += 1

    document.save(OUTPUT_DOCX)
    print(f"已生成：{OUTPUT_DOCX}")


if __name__ == "__main__":
    build_docx()
